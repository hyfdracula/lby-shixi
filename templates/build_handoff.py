"""第一阶段产出 handoff 数据包(图+数据+专业文本), 供第二阶段 CC 自写 report_draft.md。

读 config + outputs/report_stats.json + outputs/ 图, 填 handoff/ 模板占位, copy 图到 figures/。
**不出 PDF/Word** — 第一阶段只产数据包, 第二阶段由 CC 读 handoff 自写(已内化深度写作流程, 不依赖外部 skill)。

用法: python build_handoff.py -c config.yaml
产出: handoff/ (figures/ + data/stats.json + methods/study_area/results_summary/core_claims/captions + handoff_meta.json)
"""
from __future__ import annotations

import argparse
import json
import shutil
from pathlib import Path
from typing import Any

import yaml


def load_config(p):
    with open(p, encoding="utf-8") as f:
        return yaml.safe_load(f)


def _fmt_dict(d, suffix="%"):
    """dict → 'k1 v1%; k2 v2%' 字符串(供 md 占位)。"""
    return "; ".join(f"{k} {v}{suffix}" for k, v in d.items()) if d else "—"


def _build_cover(cfg, handoff):
    """生成封面 docx: 北林校徽(~84mm) + 24pt标题加粗 + 信息栏居中 + 日期
    中文字体经 w:eastAsia 显式设置, 避免回落默认字体(与 report_docx_builder._add_cover 对齐)。"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from datetime import date

    def _ea(run, font):
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.append(rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            rFonts.set(qn(attr), font)

    course = cfg.get("course", {})
    doc = Document()
    logo = Path(__file__).resolve().parent / "report" / "assets" / "bilin_logo.png"
    if logo.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Cm(8.4))
    for text in ["基于遥感植被参数的植被物候期提取", f"——以{cfg['roi']['name']}为例"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.bold = True; r.font.size = Pt(24); _ea(r, "楷体")
    for label, val in [("学　　院", course.get("college", "林学院")),
                       ("班　　级", course.get("class_name", "地信23")),
                       ("学　　号", course.get("student_id", "")),
                       ("姓　　名", course.get("student_name", "")),
                       ("指导教师", course.get("instructor", "梁博毅"))]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}：{val}"); r.font.size = Pt(14); _ea(r, "宋体")
    t = date.today()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{t.year}年{t.month}月{t.day}日"); r.font.size = Pt(12); _ea(r, "宋体")
    doc.save(str(handoff / "cover.docx"))



def main():
    ap = argparse.ArgumentParser(description="第一阶段: 产出 handoff 数据包供写作 skill")
    ap.add_argument("-c", "--config", default="config.yaml")
    args = ap.parse_args()
    cfg = load_config(args.config)
    project_root = Path(__file__).resolve().parent
    out = Path(cfg["paths"]["outputs"])
    if not out.is_absolute():
        out = project_root / out
    handoff = Path(cfg["paths"].get("handoff", "./handoff"))
    if not handoff.is_absolute():
        handoff = project_root / handoff
    handoff.mkdir(parents=True, exist_ok=True)

    # ── stats ──
    stats_path = out / "report_stats.json"
    stats: dict[str, Any] = {}
    if stats_path.exists():
        stats = json.loads(stats_path.read_text(encoding="utf-8"))

    # ── copy 图 → figures/ ──
    figs = handoff / "figures"
    figs.mkdir(exist_ok=True)
    n_fig = 0
    for sub in ["ndvi_trend", "phenology_trend", "phenology", "multi", "geo", "extra"]:
        d = out / sub
        if d.exists():
            for f in d.glob("*.png"):
                shutil.copy2(f, figs / f.name)
                n_fig += 1

    # ── data/ ──
    data_dir = handoff / "data"
    data_dir.mkdir(exist_ok=True)
    if stats_path.exists():
        shutil.copy2(stats_path, data_dir / "stats.json")

    # ── 占位填充 ──
    s = int(str(cfg["date"]["start"])[:4])
    e = int(str(cfg["date"]["end"])[:4])
    ndvi_pct = stats.get("ndvi_trend_pct", {})
    sos = stats.get("sos_mean_doy")
    eos = stats.get("eos_mean_doy")
    grow = int(eos - sos) if (sos and eos) else 230
    replacements = {
        "{研究区}": str(cfg["roi"]["name"]),
        "{时段}": f"{cfg['date']['start']} 至 {cfg['date']['end']}",
        "{年数}": str(e - s + 1),
        "{500m/降采样}": f"{cfg.get('data', {}).get('scale', 500)}m",
        "{w}": str(cfg.get("preprocess", {}).get("sg_window", 5)),
        "{p}": str(cfg.get("preprocess", {}).get("sg_polyorder", 2)),
        "{threshold_ratio}": str(cfg.get("phenology", {}).get("threshold_ratio", 0.2)),
        "{evergreen_thr}": str(cfg.get("phenology", {}).get("evergreen_thr", 0.1)),
        "{ndvi_slope_thr}": str(cfg.get("trend", {}).get("slope_threshold", 0.0005)),
        "{pheno_slope_thr}": str(cfg.get("trend", {}).get("pheno_slope_threshold", 0.5)),
        "{ndvi_slope_mean}": str(stats.get("ndvi_slope_mean", "—")),
        "{ndvi_trend_pct}": _fmt_dict(ndvi_pct),
        "{ndvi_slope_by_veg}": _fmt_dict(stats.get("ndvi_slope_by_veg", {}), ""),
        "{lai_slope}": str(stats.get("lai_slope", "—")),
        "{gpp_slope}": str(stats.get("gpp_slope", "—")),
        "{lai_mean}": str(stats.get("lai_mean", "—")),
        "{gpp_mean}": str(stats.get("gpp_mean", "—")),
        "{sos_mean_doy}": str(stats.get("sos_mean_doy", "—")),
        "{peak_mean_doy}": str(stats.get("peak_mean_doy", "—")),
        "{eos_mean_doy}": str(stats.get("eos_mean_doy", "—")),
        "{sos_trend_pct}": _fmt_dict(stats.get("sos_trend_pct", {})),
        "{eos_trend_pct}": _fmt_dict(stats.get("eos_trend_pct", {})),
        "{peak_trend_pct}": _fmt_dict(stats.get("peak_trend_pct", {})),
        "{显著改善占比}": str(ndvi_pct.get("显著改善", "—")),
        "{生长季}": str(grow),
        "{气候带}": "{【气候带, CC 填】}",  # 留给 study_area 的 CC 填
    }

    # ── 填 md 模板 ──
    tmpl = project_root / "handoff"
    for md in ["methods.md", "study_area.md", "results_summary.md", "core_claims.md", "captions.md", "code_snippets.md"]:
        src = tmpl / md
        if not src.exists():
            continue
        t = src.read_text(encoding="utf-8")
        for k, v in replacements.items():
            t = t.replace(k, v)
        (handoff / md).write_text(t, encoding="utf-8")

    # ── handoff_meta.json (text replace, 字符串值可接受) ──
    meta_src = tmpl / "handoff_meta.json"
    if meta_src.exists():
        meta = meta_src.read_text(encoding="utf-8")
        for k, v in replacements.items():
            meta = meta.replace(k, str(v))
        (handoff / "handoff_meta.json").write_text(meta, encoding="utf-8")

    # ── 封面 cover.docx (北林校徽+24pt标题+信息栏+日期) ──
    try:
        _build_cover(cfg, handoff)
        print("   cover.docx (封面)")
    except Exception as e:  # noqa: BLE001
        print(f"   ⚠ cover.docx 生成失败: {e}")

    # ── 残留占位检查 ──
    import re as _re
    _ph = _re.compile(r"\{[a-z_][a-z0-9_]*\}")
    ph_residual = 0
    bracket_residual = 0
    for md in ["study_area.md", "core_claims.md", "results_summary.md", "captions.md", "methods.md"]:
        p = handoff / md
        if not p.exists():
            continue
        for ln in p.read_text(encoding="utf-8").splitlines():
            if _ph.search(ln):
                ph_residual += 1
            if "【" in ln:
                bracket_residual += 1

    print(f"✅ handoff 数据包产出 -> {handoff}")
    print(f"   figures/ ({n_fig} 张图) + data/stats.json")
    print(f"   methods.md / study_area.md / results_summary.md / core_claims.md / captions.md / handoff_meta.json")
    if ph_residual:
        print(f"   ⚠️ 仍有 {ph_residual} 行 {{xxx}} 模板占位未回填 (stats.json 缺该字段?)")
    if bracket_residual:
        print(f"   ⚠️ 仍有 {bracket_residual} 行【CC填】占位(研究区地理/气候/生态工程), 手填后再给写作 skill")
    print("\n📋 第二阶段: CC 按 references/stage2-writing-guide.md 自写 report_draft.md")



if __name__ == "__main__":
    main()
