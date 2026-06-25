from __future__ import annotations

import argparse
import json
import logging
import re
import shutil
from pathlib import Path
from typing import Iterable

import yaml
from PIL import Image
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


logger = logging.getLogger(__name__)

CONTENT_WIDTH_IN = 6.25
BODY_FONT = "宋体"
HEADING_FONT = "黑体"
COVER_FONT = "楷体"
CODE_FONT = "Consolas"
HEADING_COLOR = "1F4D78"

VEGETATION_TYPE_PHENOLOGY_FIGURES = [
    "veg_sos_yearly",
    "veg_eos_yearly",
    "veg_peak_yearly",
    "SOS_veg_bar",
    "EOS_veg_bar",
    "Peak_veg_bar",
    "veg_SOS_trend_stacked",
    "veg_EOS_trend_stacked",
    "veg_Peak_trend_stacked",
]


SECTION_FIGURES = {
    "2.3.2": ["lc_reclass"],
    "2.3.3": ["raw_vs_smooth"],
    "2.3.4": ["phenology_diagram"],
    "4.1.1": ["ndvi_yearly", "trend_spatial_geo", "trend_pie", "trend_spatial"],
    "4.1.2": ["sen_by_veg", "veg_ndvi_yearly", "veg_ndvi_trend_stacked"],
    "4.2.1": [
        "sos_mean_geo",
        "peak_mean_geo",
        "eos_mean_geo",
        "sos_mean",
        "peak_mean",
        "eos_mean",
    ],
    "4.2.2": [
        "pheno_yearly",
        "SOS_pie",
        "SOS_spatial_geo",
        "SOS_spatial",
        "EOS_pie",
        "EOS_spatial_geo",
        "EOS_spatial",
        "Peak_pie",
        "Peak_spatial_geo",
        "Peak_spatial",
        *VEGETATION_TYPE_PHENOLOGY_FIGURES,
    ],
    "4.3": ["lai_yearly", "gpp_yearly", "multi_compare"],
}


APPENDIX_GROUPS = [
    (
        "NDVI趋势与植被类型",
        [
            "trend_spatial_geo",
            "trend_spatial",
            "trend_pie",
            "ndvi_yearly",
            "sen_by_veg",
            "veg_ndvi_yearly",
            "veg_ndvi_trend_stacked",
        ],
    ),
    (
        "物候均值与趋势",
        [
            "sos_mean_geo",
            "sos_mean",
            "peak_mean_geo",
            "peak_mean",
            "eos_mean_geo",
            "eos_mean",
            "pheno_yearly",
            "SOS_pie",
            "SOS_spatial_geo",
            "SOS_spatial",
            "EOS_pie",
            "EOS_spatial_geo",
            "EOS_spatial",
            "Peak_pie",
            "Peak_spatial_geo",
            "Peak_spatial",
        ],
    ),
    (
        "按植被类型统计",
        VEGETATION_TYPE_PHENOLOGY_FIGURES,
    ),
    (
        "多源与方法辅助图",
        [
            "lai_yearly",
            "gpp_yearly",
            "multi_compare",
            "raw_vs_smooth",
            "phenology_diagram",
            "lc_reclass",
        ],
    ),
]


FALLBACK_CAPTIONS = {
    "trend_spatial_geo": "{roi}NDVI趋势空间分布（地理坐标与研究区边界）。",
    "trend_spatial": "{roi}NDVI趋势空间分布（分类栅格视角）。",
    "trend_pie": "{roi}NDVI趋势等级面积占比。",
    "ndvi_yearly": "{roi}区域年均NDVI年际变化。",
    "sen_by_veg": "不同植被类型NDVI Sen斜率对比。",
    "veg_ndvi_yearly": "不同植被类型NDVI年际变化。",
    "veg_ndvi_trend_stacked": "不同植被类型NDVI趋势等级占比。",
    "sos_mean_geo": "SOS多年均值空间分布（地理坐标）。",
    "sos_mean": "SOS多年均值空间分布。",
    "peak_mean_geo": "Peak多年均值空间分布（地理坐标）。",
    "peak_mean": "Peak多年均值空间分布。",
    "eos_mean_geo": "EOS多年均值空间分布（地理坐标）。",
    "eos_mean": "EOS多年均值空间分布。",
    "pheno_yearly": "SOS、Peak与EOS区域均值年际变化。",
    "SOS_pie": "SOS趋势等级面积占比。",
    "SOS_spatial_geo": "SOS趋势空间分布（地理坐标）。",
    "SOS_spatial": "SOS趋势空间分布。",
    "EOS_pie": "EOS趋势等级面积占比。",
    "EOS_spatial_geo": "EOS趋势空间分布（地理坐标）。",
    "EOS_spatial": "EOS趋势空间分布。",
    "Peak_pie": "Peak趋势等级面积占比。",
    "Peak_spatial_geo": "Peak趋势空间分布（地理坐标）。",
    "Peak_spatial": "Peak趋势空间分布。",
    "veg_sos_yearly": "不同植被类型SOS年际变化。",
    "veg_eos_yearly": "不同植被类型EOS年际变化。",
    "veg_peak_yearly": "不同植被类型Peak年际变化。",
    "SOS_veg_bar": "不同植被类型SOS统计对比。",
    "EOS_veg_bar": "不同植被类型EOS统计对比。",
    "Peak_veg_bar": "不同植被类型Peak统计对比。",
    "veg_SOS_trend_stacked": "不同植被类型SOS趋势等级占比。",
    "veg_EOS_trend_stacked": "不同植被类型EOS趋势等级占比。",
    "veg_Peak_trend_stacked": "不同植被类型Peak趋势等级占比。",
    "lai_yearly": "LAI区域年均值年际变化。",
    "gpp_yearly": "GPP区域年均值年际变化。",
    "multi_compare": "NDVI、LAI与GPP多源归一化趋势对比。",
    "raw_vs_smooth": "原始NDVI序列与SG滤波后序列对比。",
    "phenology_diagram": "动态阈值法提取SOS、Peak与EOS的示意图。",
    "lc_reclass": "IGBP植被类型重分类结果。",
}


def _get_or_add(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_font(obj, font: str) -> None:
    if hasattr(obj, "font"):
        obj.font.name = font
    element = getattr(obj, "_element", None)
    if element is None:
        return
    rpr = element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font)


def _format_run(run, font: str = BODY_FONT, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = font
    _set_font(run, font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _configure_styles(doc: Document) -> None:
    """Configure fixed document styles: body 宋体, headings 黑体, code Consolas."""
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    _set_font(normal, BODY_FONT)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in [
        ("Heading 1", 16, 14, 8),
        ("Heading 2", 14, 10, 6),
        ("Heading 3", 12.5, 8, 4),
    ]:
        style = styles[name]
        style.font.name = HEADING_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(HEADING_COLOR)
        _set_font(style, HEADING_FONT)
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = BODY_FONT
    caption.font.size = Pt(10)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(80, 80, 80)
    _set_font(caption, BODY_FONT)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = None
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_together = True

    if "Code Block" not in styles:
        styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code = styles["Code Block"]
    code.font.name = CODE_FONT
    code.font.size = Pt(8.5)
    code.paragraph_format.first_line_indent = None
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)


def _shade_paragraph(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_table_borders(table, val: str = "nil") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), "0")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "FFFFFF")


def _set_table_geometry(table, widths_in: list[float]) -> None:
    widths_dxa = [int(round(width * 1440)) for width in widths_in]
    total = sum(widths_dxa)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = _get_or_add(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    tbl_ind = _get_or_add(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    layout = _get_or_add(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = _get_or_add(tc_pr, "w:tcW")
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            _set_cell_margins(cell)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def _load_config(project_root: Path) -> dict:
    cfg_path = project_root / "config.yaml"
    if not cfg_path.exists():
        return {}
    return yaml.safe_load(cfg_path.read_text(encoding="utf-8")) or {}


def _load_stats(project_root: Path) -> dict:
    path = project_root / "outputs" / "report_stats.json"
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def _count_chinese_chars(text: str) -> int:
    """Count Chinese characters in a string (Unicode range U+4E00–U+9FFF)."""
    return sum(1 for ch in text if "一" <= ch <= "鿿")


def _extract_study_area_text(project_root: Path, max_chars: int = 1200) -> str:
    """Extract full study_area.md text, truncated to max_chars Chinese characters."""
    path = project_root / "handoff" / "study_area.md"
    if not path.exists():
        return "{roi}地处研究区，属典型气候带，植被类型多样。"
    raw = path.read_text(encoding="utf-8")
    lines = []
    for line in raw.splitlines():
        s = line.strip()
        if not s or s.startswith("#") or s.startswith(">"):
            continue
        lines.append(s)
    text = "\n\n".join(lines)
    if _count_chinese_chars(text) <= max_chars:
        return text
    # Truncate to max_chars Chinese chars
    result: list[str] = []
    count = 0
    for ch in text:
        result.append(ch)
        if "一" <= ch <= "鿿":
            count += 1
            if count >= max_chars:
                break
    return "".join(result)


def _add_cover(doc: Document, project_root: Path, cfg: dict) -> None:
    """Course-style cover from cover_template.docx, replacing name/student_id/roi.

    Uses the embedded cover_template.docx (李奕葳's original, 35KB) as the base,
    replaces: 学号(210302221→config), 姓名(李奕葳→config), 班级(地信22→config),
    研究区(湖南→config.roi.name). Logo, title font, spacing all preserved from template.
    """
    from copy import deepcopy

    course = cfg.get("course", {})
    roi_name = cfg.get("roi", {}).get("name", "研究区")

    # 1. Find cover_template.docx
    template_candidates = [
        Path(__file__).resolve().parent / "assets" / "cover_template.docx",
        project_root / "report" / "assets" / "cover_template.docx",
        Path.home() / ".claude" / "skills" / "lby-shixi" / "templates" / "assets" / "cover_template.docx",
    ]
    cover_path = None
    for c in template_candidates:
        if c.exists():
            cover_path = c
            break

    if cover_path is None:
        # Fallback: build cover from scratch (old method)
        _add_cover_fallback(doc, project_root, cfg)
        return

    # 2. Load cover_template.docx and copy its body into doc
    cover_doc = Document(str(cover_path))

    # Replacements: template value → config value
    replacements = {
        "210302221": str(course.get("student_id", "")),
        "李奕葳": course.get("student_name", ""),
        "地信 22": course.get("class_name", "地信23"),
        "地信22": course.get("class_name", "地信23"),
        "湖南": roi_name,
    }

    # 3. Copy paragraphs from cover_doc body into doc, applying replacements
    for para in cover_doc.paragraphs:
        new_para = deepcopy(para._element)
        # Apply text replacements in all runs
        for run_elem in new_para.findall(qn('w:r')):
            for t_elem in run_elem.findall(qn('w:t')):
                if t_elem.text:
                    for old_val, new_val in replacements.items():
                        t_elem.text = t_elem.text.replace(old_val, new_val)
        doc.element.body.insert(list(doc.element.body).index(doc.element.body[-1]), new_para)

    # 4. Copy images (inline shapes) from cover_doc
    for rel in cover_doc.part.rels.values():
        if "image" in rel.reltype:
            # Images come with the paragraph elements via deepcopy above
            pass

    # 5. Page break after cover
    doc.add_page_break()


def _add_cover_fallback(doc: Document, project_root: Path, cfg: dict) -> None:
    """Fallback: build cover from scratch when cover_template.docx not found."""
    course = cfg.get("course", {})
    roi_name = cfg.get("roi", {}).get("name", "研究区")
    logo = project_root / "report" / "assets" / "bilin_logo.png"

    def blank(count: int, line_spacing: float = 1.0) -> None:
        for _ in range(count):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_after = Pt(0)

    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(logo), width=Cm(8.8))

    blank(2, 1.12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = None
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run("基于遥感植被参数的植被物候期提取")
    _format_run(run, COVER_FONT, 24, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = None
    subtitle.paragraph_format.space_after = Pt(0)
    run = subtitle.add_run(f"——以{roi_name}为例")
    _format_run(run, COVER_FONT, 24, True)

    blank(6, 1.1)

    fields = [
        ("学    院：", course.get("college", "")),
        ("班    级：", course.get("class_name", "")),
        ("学    号：", str(course.get("student_id", ""))),
        ("姓    名：", course.get("student_name", "")),
        ("指导教师：", course.get("instructor", "")),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 0.92
        p.paragraph_format.space_after = Pt(3)
        left = p.add_run(label)
        gap = "    " if label != "指导教师：" else "   "
        right = p.add_run(f"{gap}{value}    ")
        _format_run(left, COVER_FONT, 15)
        value_font = "Times New Roman" if str(value).isascii() else COVER_FONT
        _format_run(right, value_font, 15)
        right.underline = True

    date_text = cfg.get("course", {}).get("date") or "2026年6月24日"
    date_text = re.sub(r"(\d{4})年\s*(\d{1,2})月\s*(\d{1,2})日", r"\1 年 \2 月 \3 日", date_text)
    blank(2, 1.18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(date_text)
    _format_run(run, COVER_FONT, 10.5)

    doc.add_page_break()


def _body_lines(markdown_path: Path) -> list[str]:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    for idx, line in enumerate(lines):
        if line.startswith("## "):
            return lines[idx:]
    return lines


def _caption_map(captions_path: Path, stats: dict) -> dict[str, str]:
    if not captions_path.exists():
        return {}
    current = ""
    mapping: dict[str, str] = {}
    for raw in captions_path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if line.startswith("**Figure"):
            current = re.sub(r"^\*\*Figure\s+\d+\.\*\*\s*", "", line).strip()
            current = _clean_caption(current, stats)
        elif "图文件：" in line and current:
            match = re.search(r"figures/([^)*]+\.png)", line)
            if match:
                mapping[match.group(1)] = current
    return mapping


def _clean_caption(text: str, stats: dict | None = None) -> str:
    text = re.sub(r"^\s*(图|Figure)\s*\d+\s*[.．、:：]?\s*", "", text.strip())
    text = text.replace("Figure", "图")
    if stats:
        text = text.replace("{lai_slope}", str(stats.get("lai_slope", "见统计结果")))
        text = text.replace("{gpp_slope}", str(stats.get("gpp_slope", "见统计结果")))
    text = re.sub(r"\{[^}]+\}", "见统计结果", text)
    return text.strip()


def _fallback_caption(stem: str, roi_name: str) -> str:
    return FALLBACK_CAPTIONS.get(stem, stem.replace("_", " ")).format(roi=roi_name)


def _normalise_figure_caption(
    basename: str,
    provided_caption: str | None,
    captions: dict[str, str],
    roi_name: str,
    stats: dict,
) -> str:
    if provided_caption:
        caption = provided_caption
    else:
        caption = captions.get(basename) or _fallback_caption(Path(basename).stem, roi_name)
    return _clean_caption(caption, stats)


def _resolve_image(project_root: Path, markdown_path: Path, raw_path: str) -> Path | None:
    cleaned = raw_path.strip().replace("\\", "/")
    if re.match(r"^[a-z]+://", cleaned):
        return None
    candidates = [
        markdown_path.parent / cleaned,
        project_root / cleaned,
        project_root / cleaned.lstrip("./"),
    ]
    for candidate in candidates:
        try:
            resolved = candidate.resolve()
        except OSError:
            continue
        if resolved.exists():
            return resolved
    return None


def _add_heading(doc: Document, text: str, markdown_level: int) -> None:
    level = min(max(markdown_level - 1, 1), 3)
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.first_line_indent = None
    for run in heading.runs:
        _format_run(run, HEADING_FONT, None, True)


def _add_text_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _format_run(run, BODY_FONT, 12, True)
        else:
            run = paragraph.add_run(part)
            _format_run(run, BODY_FONT, 12)


def _split_table_row(line: str) -> list[str]:
    _PH = "\x00"
    cleaned = line.strip().strip("|").replace(r"\|", _PH)
    return [cell.replace(_PH, "|").strip() for cell in cleaned.split("|")]


def _is_separator_row(cells: Iterable[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def _column_widths(headers: list[str]) -> list[float]:
    cols = len(headers)
    if cols == 5 and any("产品" in item for item in headers):
        return [0.9, 2.15, 1.05, 1.15, 1.25]
    if cols == 6 and any("物候" in item for item in headers):
        return [0.85, 1.05, 1.05, 1.35, 1.1, 1.1]
    return [CONTENT_WIDTH_IN / cols] * cols


def _add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = [_split_table_row(line) for line in lines if line.strip()]
    rows = [row for row in rows if not _is_separator_row(row)]
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = _column_widths(rows[0])
    _set_table_geometry(table, widths)

    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                if c_idx > 1 or len(row) <= 3:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    _format_run(run, BODY_FONT, 10.5, r_idx == 0)
            if r_idx == 0:
                _shade_cell(cell, "E8EEF5")
    _repeat_header(table.rows[0])
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    spacer.paragraph_format.first_line_indent = None


def _add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph(style="Code Block")
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.keep_together = True
        _shade_paragraph(paragraph, "F4F6F9")
        run = paragraph.add_run(line if line else " ")
        _format_run(run, "Consolas", 8.5)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.first_line_indent = None
    spacer.paragraph_format.space_after = Pt(6)


def _figure_grid_rows(count: int) -> list[int]:
    # 每行最多 2 张图: 1→[1], 2→[2], 3→[2,1], 4→[2,2], 5→[2,2,1], 6→[2,2,2]...
    if count <= 0:
        return []
    if count == 1:
        return [1]
    rows = [2] * (count // 2)
    if count % 2 == 1:
        rows.append(1)
    return rows


def _image_display_width(path: Path) -> float:
    with Image.open(path) as image:
        width, height = image.size
    ratio = width / height if height else 1.0
    max_height = 5.15
    max_width = 4.7 if 0.85 <= ratio <= 1.15 else 5.9
    return min(max_width, max_height * ratio)


def _image_cell_display_width(path: Path, columns: int) -> float:
    with Image.open(path) as image:
        width, height = image.size
    ratio = width / height if height else 1.0
    cell_width = (CONTENT_WIDTH_IN / columns) - 0.16
    max_height = 2.35 if columns == 2 else 1.55
    return max(1.2, min(cell_width, max_height * ratio))


def _add_figure(
    doc: Document,
    image_path: Path,
    caption: str,
    figure_state: dict[str, int],
    seen: set[str],
) -> None:
    if not image_path.exists() or image_path.name in seen:
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(image_path), width=Inches(_image_display_width(image_path)))

    figure_state["figure"] += 1
    cap = doc.add_paragraph(style="Caption")
    cap.paragraph_format.first_line_indent = None
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"图 {figure_state['figure']}　{caption}")
    _format_run(run, BODY_FONT, 10)
    seen.add(image_path.name)


def _add_figure_group(
    doc: Document,
    figures: list[tuple[Path, str]],
    figure_state: dict[str, int],
    seen: set[str],
) -> None:
    pending = [(path, caption) for path, caption in figures if path.exists() and path.name not in seen]
    if not pending:
        return
    if len(pending) == 1:
        _add_figure(doc, pending[0][0], pending[0][1], figure_state, seen)
        return

    offset = 0
    for columns in _figure_grid_rows(len(pending)):
        row_items = pending[offset:offset + columns]
        offset += columns
        table = doc.add_table(rows=1, cols=columns)
        table.autofit = False
        _set_table_geometry(table, [CONTENT_WIDTH_IN / columns] * columns)
        _set_table_borders(table)
        for idx, (image_path, caption) in enumerate(row_items):
            cell = table.cell(0, idx)
            _set_cell_margins(cell, top=40, bottom=40, start=60, end=60)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            image_paragraph = cell.paragraphs[0]
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_paragraph.paragraph_format.first_line_indent = None
            image_paragraph.paragraph_format.keep_with_next = True
            image_paragraph.paragraph_format.space_after = Pt(1)
            image_paragraph.add_run().add_picture(
                str(image_path),
                width=Inches(_image_cell_display_width(image_path, columns)),
            )

            figure_state["figure"] += 1
            caption_paragraph = cell.add_paragraph(style="Caption")
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_paragraph.paragraph_format.first_line_indent = None
            caption_paragraph.paragraph_format.space_after = Pt(2)
            run = caption_paragraph.add_run(f"图 {figure_state['figure']}　{caption}")
            _format_run(run, BODY_FONT, 8)
            seen.add(image_path.name)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.first_line_indent = None
        spacer.paragraph_format.space_after = Pt(3)


def _figures_for_heading(heading: str) -> list[str]:
    result: list[str] = []
    for key, names in SECTION_FIGURES.items():
        if key in heading:
            result.extend(names)
    return result


def _figure_path(project_root: Path, stem: str) -> Path:
    return project_root / "handoff" / "figures" / f"{stem}.png"


def _add_section_figures(
    doc: Document,
    project_root: Path,
    heading: str | None,
    captions: dict[str, str],
    roi_name: str,
    stats: dict,
    figure_state: dict[str, int],
    seen: set[str],
) -> None:
    if not heading:
        return

    stems = _figures_for_heading(heading)

    def collect(stem_names: list[str]) -> list[tuple[Path, str]]:
        figures = []
        for stem in stem_names:
            path = _figure_path(project_root, stem)
            caption = _normalise_figure_caption(path.name, None, captions, roi_name, stats)
            figures.append((path, caption))
        return figures

    if "4.2.2" in heading:
        vegetation_stems = [stem for stem in stems if stem in VEGETATION_TYPE_PHENOLOGY_FIGURES]
        main_stems = [stem for stem in stems if stem not in VEGETATION_TYPE_PHENOLOGY_FIGURES]
        _add_figure_group(doc, collect(main_stems), figure_state, seen)
        pending_vegetation = [
            stem for stem in vegetation_stems if _figure_path(project_root, stem).exists()
            and _figure_path(project_root, stem).name not in seen
        ]
        if pending_vegetation:
            _add_text_paragraph(
                doc,
                "**按植被类型统计：** 为进一步对比不同植被类型对物候变化的响应，补充展示SOS、EOS与Peak的年际曲线、类型统计和趋势等级占比。",
            )
            _add_figure_group(doc, collect(pending_vegetation), figure_state, seen)
        return

    figures = []
    for stem in stems:
        path = _figure_path(project_root, stem)
        caption = _normalise_figure_caption(path.name, None, captions, roi_name, stats)
        figures.append((path, caption))
    _add_figure_group(doc, figures, figure_state, seen)


def _add_appendix_figures(
    doc: Document,
    project_root: Path,
    captions: dict[str, str],
    roi_name: str,
    stats: dict,
    figure_state: dict[str, int],
    seen: set[str],
) -> None:
    figures_dir = project_root / "handoff" / "figures"
    remaining = {path.stem for path in figures_dir.glob("*.png") if path.name not in seen}

    # 统一标题：附录A 输出图件汇总（无论 remaining 是否为空都生成）
    doc.add_page_break()
    _add_heading(doc, "附录A 输出图件汇总", 2)

    if remaining:
        for group_title, names in APPENDIX_GROUPS:
            group_remaining = [name for name in names if name in remaining]
            if not group_remaining:
                continue
            _add_heading(doc, group_title, 3)
            figures = []
            for stem in group_remaining:
                path = _figure_path(project_root, stem)
                caption = _normalise_figure_caption(path.name, None, captions, roi_name, stats)
                figures.append((path, caption))
            _add_figure_group(doc, figures, figure_state, seen)
            for stem in group_remaining:
                remaining.discard(stem)

        figures = []
        for stem in sorted(remaining):
            path = _figure_path(project_root, stem)
            caption = _normalise_figure_caption(path.name, None, captions, roi_name, stats)
            figures.append((path, caption))
        _add_figure_group(doc, figures, figure_state, seen)

    # ── 附录A 子节：图件清单表格（原附B，合并入附录A）──
    _add_heading(doc, "图件清单", 3)
    all_figures = sorted(
        [path for path in (project_root / "handoff" / "figures").glob("*.png")],
        key=lambda p: p.stem,
    )
    summary_table = doc.add_table(rows=len(all_figures) + 1, cols=3)
    summary_table.style = "Table Grid"
    _set_table_geometry(summary_table, [1.0, 3.0, 2.5])
    # 表头
    for ci, header in enumerate(["序号", "文件名", "内容"]):
        cell = summary_table.cell(0, ci)
        cell.text = header
        for p in cell.paragraphs:
            p.paragraph_format.first_line_indent = None
            for run in p.runs:
                run.bold = True
                _format_run(run, BODY_FONT, 9)
    # 数据行
    for ri, path in enumerate(all_figures, start=1):
        caption = _normalise_figure_caption(
            path.name, None, captions, roi_name, stats,
        )
        for ci, text in enumerate([str(ri), path.name, caption]):
            cell = summary_table.cell(ri, ci)
            cell.text = text
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    _format_run(run, BODY_FONT, 8)
    logger.info("附录A 图件清单: %d 张图", len(all_figures))


def _add_core_code_from_handoff(doc: Document, project_root: Path) -> bool:
    path = project_root / "handoff" / "code_snippets.md"
    if not path.exists():
        return False

    lines = path.read_text(encoding="utf-8").splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or line.startswith("# 核心代码段落") or line.startswith(">"):
            idx += 1
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            _add_heading(doc, heading_match.group(2).strip(), len(heading_match.group(1)) + 1)
            idx += 1
            continue

        if line.startswith("|") and line.count("|") >= 2:
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            _add_markdown_table(doc, table_lines)
            continue

        if line.startswith("```"):
            idx += 1
            code_lines = []
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx].rstrip())
                idx += 1
            if idx < len(lines):
                idx += 1
            _add_code_block(doc, code_lines)
            continue

        _add_text_paragraph(doc, line)
        idx += 1
    return True


def _skip_to_next_top_heading(lines: list[str], idx: int) -> int:
    while idx < len(lines):
        if re.match(r"^##\s+", lines[idx].strip()):
            return idx
        idx += 1
    return idx


def _collect_image_block(
    project_root: Path,
    markdown_path: Path,
    lines: list[str],
    idx: int,
    captions: dict[str, str],
    roi_name: str,
    stats: dict,
) -> tuple[list[tuple[Path, str]], int]:
    figures: list[tuple[Path, str]] = []
    current = idx
    while current < len(lines):
        line = lines[current].strip()
        match = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", line)
        if not match:
            break
        image_path = _resolve_image(project_root, markdown_path, match.group(2))
        if image_path:
            caption = _normalise_figure_caption(
                image_path.name,
                match.group(1),
                captions,
                roi_name,
                stats,
            )
            figures.append((image_path, caption))

        lookahead = current + 1
        while lookahead < len(lines) and not lines[lookahead].strip():
            lookahead += 1
        if lookahead < len(lines) and re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", lines[lookahead].strip()):
            current = lookahead
        else:
            return figures, current + 1
    return figures, current


def build_report_docx(
    project_root: Path | str = Path("."),
    markdown_path: Path | str | None = None,
    output_path: Path | str | None = None,
    include_appendix: bool = True,
) -> Path:
    project_root = Path(project_root).resolve()
    cfg = _load_config(project_root)
    stats = _load_stats(project_root)

    if markdown_path is not None:
        # Path explicitly provided — use it directly (caller already resolved)
        markdown_path = Path(markdown_path)
        if not markdown_path.is_absolute():
            markdown_path = project_root / markdown_path
        if not markdown_path.exists():
            raise FileNotFoundError(f"Markdown draft not found: {markdown_path}")
    else:
        # Default: use CC 自写 report_draft.md
        markdown_path = project_root / "paper_rewriting_output" / "report_draft.md"
        if not markdown_path.exists():
            raise FileNotFoundError(f"Markdown draft not found: {markdown_path}")

    roi_name = cfg.get("roi", {}).get("name", "研究区")
    student_name = cfg.get("course", {}).get("student_name", "学生")
    captions = _caption_map(project_root / "handoff" / "captions.md", stats)

    if output_path is None:
        output_path = project_root / f"{roi_name}植被物候实习报告_{student_name}.docx"
    output_path = Path(output_path)
    if not output_path.is_absolute():
        output_path = project_root / output_path

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    _configure_styles(doc)
    _add_cover(doc, project_root, cfg)

    lines = _body_lines(markdown_path)
    current_heading: str | None = None
    figure_state = {"figure": 0}
    seen_figures: set[str] = set()
    idx = 0
    while idx < len(lines):
        raw = lines[idx]
        line = raw.strip()
        if not line:
            idx += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            _add_section_figures(
                doc, project_root, current_heading, captions, roi_name, stats, figure_state, seen_figures
            )
            level = len(heading_match.group(1))
            current_heading = heading_match.group(2).strip()
            _add_heading(doc, current_heading, level)
            idx += 1
            if "核心代码" in current_heading and _add_core_code_from_handoff(doc, project_root):
                idx = _skip_to_next_top_heading(lines, idx)
            continue

        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", line)
        if image_match:
            figures, idx = _collect_image_block(project_root, markdown_path, lines, idx, captions, roi_name, stats)
            _add_figure_group(doc, figures, figure_state, seen_figures)
            continue

        if line.startswith("|") and line.count("|") >= 2:
            table_lines = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            _add_markdown_table(doc, table_lines)
            continue

        if line.startswith("```"):
            idx += 1
            code_lines = []
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx].rstrip())
                idx += 1
            if idx < len(lines):
                idx += 1
            _add_code_block(doc, code_lines)
            continue

        if line.startswith("<") and line.endswith(">"):
            idx += 1
            continue

        _add_text_paragraph(doc, line)
        idx += 1

    _add_section_figures(
        doc, project_root, current_heading, captions, roi_name, stats, figure_state, seen_figures
    )
    if include_appendix:
        _add_appendix_figures(doc, project_root, captions, roi_name, stats, figure_state, seen_figures)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(name)s: %(message)s")
    parser = argparse.ArgumentParser(
        description="Build a report (docx/md) from final template or CC 自写 Markdown.",
    )
    parser.add_argument("--project-root", default=".", help="Project root containing config.yaml and handoff/.")
    parser.add_argument("--markdown", default=None, help="Markdown draft path. Defaults to paper_rewriting_output/report_draft.md.")
    parser.add_argument("--output", default=None, help="Output path (extension derives from format).")
    parser.add_argument("--no-appendix", action="store_true", help="Do not append remaining figures.")
    parser.add_argument(
        "--format",
        default=None,
        choices=["docx", "md"],
        help="Output format (docx | md).",
    )
    args = parser.parse_args()

    project_root = Path(args.project_root)
    cfg = _load_config(project_root)
    stats = _load_stats(project_root)
    fmt = (args.format or str(cfg.get("report", {}).get("format", "docx") or "docx")).strip().lower()
    if fmt not in {"docx", "md"}:
        raise ValueError(f"Unsupported report.format='{fmt}'. Supported: docx | md.")

    # 直接用 CC 第二阶段手写的 report_draft.md（不存在则报错，不再有模板填充捷径）
    markdown_path = Path(args.markdown) if args.markdown else project_root / "paper_rewriting_output" / "report_draft.md"
    if not markdown_path.is_absolute():
        markdown_path = project_root / markdown_path
    if not markdown_path.exists():
        raise FileNotFoundError(f"Markdown draft not found: {markdown_path}")

    if fmt == "md":
        # md output: copy the resolved draft to output path
        out = Path(args.output) if args.output else project_root / "paper_rewriting_output" / "report_final.md"
        if not out.is_absolute():
            out = project_root / out
        if out.resolve() != markdown_path.resolve():
            out.parent.mkdir(parents=True, exist_ok=True)
            shutil.copyfile(markdown_path, out)
        print(out)
        return

    docx_out = build_report_docx(
        project_root=project_root,
        markdown_path=markdown_path,
        output_path=Path(args.output) if args.output else None,
        include_appendix=not args.no_appendix,
    )
    print(docx_out)


if __name__ == "__main__":
    main()
